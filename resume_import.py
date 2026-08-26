from __future__ import annotations

from io import BytesIO
from pathlib import Path
import os
import uuid

from docx import Document


ALLOWED_EXTENSIONS = {".docx", ".txt"}


def validate_resume_upload(filename: str, file_bytes: bytes) -> None:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValueError("仅支持 DOCX 和 TXT 文件。")
    max_bytes = max(1, int(os.getenv("MAX_UPLOAD_MB", "10"))) * 1024 * 1024
    if len(file_bytes) > max_bytes:
        raise ValueError(f"文件超过 {os.getenv('MAX_UPLOAD_MB', '10')} MB 上传限制。")


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    validate_resume_upload(filename, file_bytes)
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValueError("Demo V0.1 仅支持 DOCX 和 TXT 文件。")
    if suffix == ".txt":
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return file_bytes.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        raise ValueError("无法识别 TXT 文件编码。")

    document = Document(BytesIO(file_bytes))
    blocks: list[str] = []
    blocks.extend(p.text.strip() for p in document.paragraphs if p.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                blocks.append(line)
    return "\n".join(blocks).strip()


def save_original_file(user_id: str, resume_id: str, filename: str, file_bytes: bytes) -> str:
    validate_resume_upload(filename, file_bytes)
    # user_id is always the canonical ID resolved by the server identity layer.
    if user_id != "dev_user":
        user_id = str(uuid.UUID(user_id))
    upload_dir = Path("data/uploads") / user_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    safe_suffix = Path(filename).suffix.lower()
    path = upload_dir / f"{resume_id}{safe_suffix}"
    path.write_bytes(file_bytes)
    return str(path)
