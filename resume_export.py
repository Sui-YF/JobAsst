from __future__ import annotations

import re
import uuid
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import database
from resume_templates import TEMPLATES


def safe_filename_part(value: str) -> str:
    value = re.sub(r"[<>:\"/\\|?*\x00-\x1f]", "_", value or "")
    value = re.sub(r"\s+", "_", value).strip("._")
    return value[:60] or "Resume"


def render_docx(content: str, template_id: str = "ats_simple") -> bytes:
    template = TEMPLATES.get(template_id, TEMPLATES["ats_simple"])
    doc = Document()
    section = doc.sections[0]
    compact = template["density"] == "compact"
    section.top_margin = section.bottom_margin = Inches(0.5 if compact else 0.65)
    section.left_margin = section.right_margin = Inches(0.65 if compact else 0.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5 if compact else 10.5)
    for index, raw in enumerate(content.splitlines()):
        line = raw.strip()
        if not line:
            doc.add_paragraph()
            continue
        if index == 0:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line)
            run.bold = True
            run.font.size = Pt(16)
        elif line.startswith(("- ", "• ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif len(line) <= 28 and not line.endswith(("。", ".", ";", "；")):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line)
            run.bold = True
            run.font.size = Pt(11.5)
        else:
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_application_docx(user_id: str, application_id: str, candidate_name: str = "Candidate") -> tuple[str, bytes]:
    workspace = database.get_application_workspace(user_id, application_id)
    if workspace["application_status"] not in {"FINAL_READY", "PREFLIGHT_READY"}:
        raise RuntimeError("只有用户确认后的Target Resume可以导出")
    content = (workspace.get("final_resume") or {}).get("content", "")
    if not content:
        raise RuntimeError("Target Resume内容为空")
    data = render_docx(content, workspace.get("template_id", "ats_simple"))
    canonical_user = user_id if user_id == database.DEV_USER_ID else str(uuid.UUID(user_id))
    export_dir = Path("data/exports") / canonical_user
    export_dir.mkdir(parents=True, exist_ok=True)
    filename = "_".join(safe_filename_part(v) for v in (
        workspace.get("company", "Company"), workspace.get("job_title", "Role"), candidate_name
    )) + ".docx"
    path = export_dir / filename
    path.write_bytes(data)
    return str(path), data
