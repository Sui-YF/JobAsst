from io import BytesIO

from docx import Document

from resume_import import extract_resume_text


def test_txt_import_utf8():
    text = extract_resume_text("resume.txt", "中文简历\nC++ Developer".encode("utf-8"))
    assert "中文简历" in text
    assert "C++ Developer" in text


def test_docx_import_paragraph_and_table():
    document = Document()
    document.add_paragraph("UE / C++ Developer")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Company"
    table.cell(0, 1).text = "Pratts"
    buffer = BytesIO()
    document.save(buffer)
    text = extract_resume_text("resume.docx", buffer.getvalue())
    assert "UE / C++ Developer" in text
    assert "Company | Pratts" in text
